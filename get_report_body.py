from get_external import get_services_report, get_voice_call_report, get_crossmod
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from const import *


# Настройка полей для документа, можно добавить доп настройки: размер документа, ориентация и т.д.
def set_margin(doc):
    sections = doc.sections
    
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)


# Возвращает номер этажа, на котором будет располагаться данная картинка
def get_picture_number_of_floor(picture):
    floor_num = picture[:picture.find('_')]
    if round(float(floor_num), 0) == float(floor_num):
        return int(floor_num)
    else:
        return float(floor_num)


# Возвращает порядковый номер скрипта картинки для функции add_picture_caption
def get_script_number_of_picture(picture):

    # 1fl_->02<-_2g.png
    script_num = picture[picture.find("_")+1:picture.find("_")+3]
    return script_num


# Возвращает отсортированный список этажей
def get_floor_list(pictures_list):
    floor_list = []
    for elem in pictures_list:
        floor_num = get_picture_number_of_floor(elem)
        if floor_num not in floor_list and floor_num is not None:
            floor_list.append(floor_num)
    return sorted(floor_list)


# Возвращает список картинок для конкретного этажа
def get_list_of_floor_pictures(floor, _all_pictures):
    pictures_list = []
    for elem in _all_pictures:
        pic_num = get_picture_number_of_floor(elem)
        if floor == pic_num:
            pictures_list.append(elem)
    return pictures_list


# Добавляем пустую строку в файл
def add_blank(doc, number_of_blanks):
    for _ in range(number_of_blanks):
        doc.add_paragraph()


# Добавление поля с данными из окна
def add_paragraph_with_external_data(doc, key, value):
    p = doc.add_paragraph()
    date_title = p.add_run(key)
    date_title.font.name = 'Times new roman'
    date_title.font.size = Pt(14)
    date_title.bold = True
    date_text = p.add_run(f"{value}")
    date_text.font.name = 'Times new roman'
    date_text.font.size = Pt(14)


# Добавляет первую страницу в отчет
def add_first_page_in_doc(doc, user_data_dict):

    # Заголовок файла
    report_title = doc.add_paragraph()
    report_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    report_title_text = report_title.add_run("Отчет по результатам измерений")
    report_title_text.font.name = 'Times new roman'
    report_title_text.font.size = Pt(20)
    report_title_text.bold = True

    # add_blank(doc, 2)

    # Добавление лого
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("")
    run.add_picture("logo.png", width=Pt(70))
    add_blank(doc, 1)

    # Данные об измерениях
    for key, value in sorted(user_data_dict.items()):
        if key == external[6] or value in [None, 'None', ''] and key not in external[3:6]:
            pass
        else:
            add_paragraph_with_external_data(doc, key[3:], value)

    # Примечание
    p = doc.add_paragraph()
    note = p.add_run("Примечание: ")
    note.font.name = 'Times new roman'
    note.font.size = Pt(14)
    note.bold = True

    # Место для картинки
    add_blank(doc, 3)
    p = doc.add_paragraph('МЕСТО ДЛЯ КАРТИНКИ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_blank(doc, 3)

    # Заключение
    # doc.add_paragraph("_" * 132)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(21)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    purpose_of_work_title = p.add_run("Заключение по результатам измерений")
    purpose_of_work_title.font.name = 'Times new roman'
    purpose_of_work_title.font.size = Pt(14)
    purpose_of_work_title.bold = True
    # purpose_of_work_text = p.add_run(f"{purpose_of_work}")
    # purpose_of_work_text.font.name = 'Times new roman'
    # purpose_of_work_text.font.size = Pt(14)

    # Текст заключения
    purpose_of_work = '\tПокрытие там-то там-то в такой-то технологии такое-то или такое-то.\n' \
                      'На таком-то этаже кое-что что-то делалось, в таком-то количестве что-то ' \
                      'обнаружили.\n' \
                      'В целом что-то как-то\n' \
                      'На какой-то БС кое-что наблюдалось.\n' \
                      'Требуется что-то, а может и не требуется.\n'
    p = doc.add_paragraph()
    result_text_of_work = p.add_run(f'{purpose_of_work}')
    result_text_of_work.font.name = 'Times new roman'
    result_text_of_work.font.size = Pt(14)
    doc.add_page_break()


# Добавление нижнего колонтитула в файл, на вход принимает документ
def add_footer_in_doc(doc, user_data_dict):
    date = user_data_dict["3. Дата измерений: "]
    building = user_data_dict["5. Объект: "]
    address = user_data_dict["4. Адрес: "]

    if date in external[:3] and \
            building in external[:3] and \
            address in external[:3]:
        pass
    else:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]

        if building is None:
            footer_text = footer_para.add_run(f'{date}\n{address}')
        else:
            footer_text = footer_para.add_run(f'{date}\n{building}\n{address}')

        footer_text.font.name = 'Times new roman'
        footer_text.font.size = Pt(8)


# Добавляет верхний колонтитул
def add_header_in_doc(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_text = header_para.add_run('Отчет по результатам проведения\n'
                                      f'измерений в сети ПАО «МегаФон»')

    header_text.font.name = 'Times new roman'
    header_text.font.size = Pt(11)
    header_text.italic = True


# Добавляет заголовок этажа и форматирует его
def add_floor_title_in_file(floor, doc):
    if OUTDOOR_FLOOR_NUMBER <= floor <= OUTDOOR_FLOOR_NUMBER + 1:
        floor_title = doc.add_heading().add_run(f'{OUTDOOR_TITLE}')
    elif UNDEFINED_FLOOR_NUMBER <= floor <= UNDEFINED_FLOOR_NUMBER + 1:
        floor_title = doc.add_heading().add_run(f'{UNDEFINED_FLOOR_TITLE} этаж')
    else:
        floor_title = doc.add_heading().add_run(f'{floor} этаж')
    floor_title.font.name = 'Times new roman'
    floor_title.font.size = Pt(16)
    floor_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


# Добавление таблицы voice call report с заголовком
def add_voice_call_report_table(path, doc):
    data = get_voice_call_report(path)
    if data == -1:
        return -1

    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Показатели голосовых вызовов')

    # Форматирование подписи
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)
    text.font.bold = True

    rows, cols = 9, 4
    vcr_table = doc.add_table(rows=rows, cols=cols)
    vcr_table.style = 'Table Grid'
    vcr_table.rows[0].cells[1].text = '2G'
    vcr_table.rows[0].cells[2].text = 'CSFB'
    vcr_table.rows[0].cells[3].text = 'VoLTE'
    i = 1
    for key in data.keys():

        vcr_table.rows[i].cells[0].text = VCR_TRANSLATE[key]
        one_kpi_data = data.get(key)
        for tech_key, val in one_kpi_data.items():
            if key in PROCENTS_VCR_KPIS:
                val += ' %'
            if tech_key == '2G':
                vcr_table.rows[i].cells[1].text = val
            elif tech_key == 'CSFB':
                vcr_table.rows[i].cells[2].text = val
            elif tech_key == 'VoLTE':
                vcr_table.rows[i].cells[3].text = val

        i += 1

    adjust_vcr_table(vcr_table, rows, cols)
    add_blank(doc, 1)
    return 0


# Добавление таблицы voice call report с заголовком
def add_services_report_table(path, doc):
    data = get_services_report(path)
    if data == -1:
        return -1

    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Сводная таблица показателей сервисов и покрытия сети по объекту в целом')

    # Форматирование подписи
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)
    text.font.bold = True

    rows = len(data.keys()) + 1
    cols = 8
    sr_table = doc.add_table(rows=rows, cols=cols)
    sr_table.style = 'Table Grid'
    for i in range(1, cols):
        sr_table.rows[0].cells[i].text = SR_COLUMNS[i]

    i = 1
    for key in sorted(data.keys()):
        one_floor_kpis = data.get(key)
        sr_table.rows[i].cells[0].text = key

        for kpi_key, val in one_floor_kpis.items():
            if kpi_key == SR_COLUMNS[1]:
                cell1 = sr_table.rows[i].cells[1]
                cell1.text = val
                set_cell_color(cell1, val, kpi_key)
            elif kpi_key == SR_COLUMNS[2]:
                cell2 = sr_table.rows[i].cells[2]
                cell2.text = val
                set_cell_color(cell2, val, kpi_key)
            elif kpi_key == SR_COLUMNS[3]:
                cell3 = sr_table.rows[i].cells[3]
                cell3.text = val
                set_cell_color(cell3, val, kpi_key)
            elif kpi_key == SR_COLUMNS[4]:
                sr_table.rows[i].cells[4].text = val
            elif kpi_key == SR_COLUMNS[5]:
                sr_table.rows[i].cells[5].text = val
            elif kpi_key == SR_COLUMNS[6]:
                sr_table.rows[i].cells[6].text = val
            elif kpi_key == SR_COLUMNS[7]:
                sr_table.rows[i].cells[7].text = val
        i += 1

    adjust_sr_table(sr_table, rows=rows, cols=cols)
    add_color_legend_table(doc)
    add_blank(doc, 1)


# Выбор цвета заливки ячейки
def set_cell_color(cell, val, technology):
    if val in [None, 'None', 'None/None']:
        return 0

    value = float(val)
    if technology == 'RxLevel':
        if value >= -70:
            shade_background_cell(cell, GREEN)
        elif -70 > value >= -85:
            shade_background_cell(cell, LIGHT_GREEN)
        elif -85 > value >= -95:
            shade_background_cell(cell, ORANGE)
        else:
            shade_background_cell(cell, RED)

    if technology == 'RSCP':
        if value >= -80:
            shade_background_cell(cell, GREEN)
        elif -80 > value >= -90:
            shade_background_cell(cell, LIGHT_GREEN)
        elif -90 > value >= -100:
            shade_background_cell(cell, ORANGE)
        else:
            shade_background_cell(cell, RED)

    if technology == 'RSRP':
        if value >= -85:
            shade_background_cell(cell, GREEN)
        elif -85 > value >= -100:
            shade_background_cell(cell, LIGHT_GREEN)
        elif -100 > value >= -110:
            shade_background_cell(cell, ORANGE)
        else:
            shade_background_cell(cell, RED)


# Заливка ячейки цветом
def shade_background_cell(cell, color="ff00ff"):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)


# Добавляет 1 картинку файл
def add_picture_in_file(_path, doc, picture):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("")
    run.add_picture(_path + '/' + picture, width=Pt(460))


# Форматирование voice call report table
def adjust_vcr_table(tab, rows, cols):
    tab.alignment = WD_TAB_ALIGNMENT.CENTER
    for cell in tab.columns[0].cells:
        cell.width = Cm(6)

    for i in range(1, cols):
        for cell in tab.columns[i].cells:
            cell.width = Cm(2)

    for i in range(rows):
        for j in range(1, cols):
            tab.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


# Форматирование services report table
def adjust_sr_table(tab, rows, cols):
    tab.alignment = WD_TAB_ALIGNMENT.CENTER
    for cell in tab.columns[0].cells:
        cell.width = Cm(4)

    for i in range(1, cols):
        for cell in tab.columns[i].cells:
            if i in [1, 2, 3]:
                cell.width = Cm(2)
            elif i in [4, 5]:
                cell.width = Cm(4)
            else:
                cell.width = Cm(3)

    for i in range(rows):
        for j in range(1, cols):
            tab.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


# Добавление таблицы с цветами
def add_color_legend_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    table.cell(0, 0).text = 'Отлично'
    shade_background_cell(table.cell(0, 0), GREEN)
    table.cell(0, 1).text = 'Хорошо'
    shade_background_cell(table.cell(0, 1), LIGHT_GREEN)
    table.cell(0, 2).text = 'Удовлетворительно'
    shade_background_cell(table.cell(0, 2), ORANGE)
    table.cell(0, 3).text = 'Неудовлетворительно'
    shade_background_cell(table.cell(0, 3), RED)

    for j in range(4):
        table.cell(0, j).paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


# Добавление информации о пересечениях PCI
def add_crossmod(path, doc):
    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Зафиксированные пересечения PCI\n')

    # Форматирование подписи
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)
    text.font.bold = True

    data = get_crossmod(path)
    if  isinstance(data, str):
        text = paragraph.add_run(f'Пересечения можно посмотреть в файле {data},'
                                 f' т.к. он слишком большой\n')

        # Форматирование подписи
        text.font.name = 'Times new roman'
        text.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return 0

    for line in data:
        paragraph = doc.add_paragraph('')
        text = paragraph.add_run(line)
        text.font.name = 'Times new roman'
        text.font.size = Pt(12)
        paragraph.paragraph_format.line_spacing = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_page_break()


# Добавление названия к картинке
def add_picture_caption(picture, doc, serial_number):

    script_number_of_picture = get_script_number_of_picture(picture)
    picture_caption = PICTURES_CAPTIONS.get(script_number_of_picture)
    paragraph = doc.add_paragraph('')
    text = paragraph.add_run(f'Рисунок {serial_number} - {picture_caption}')

    # Форматирование подписи
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.font.name = 'Times new roman'
    text.font.size = Pt(14)


# Добавление подписи для раздела измерений со сканера
def add_scanner_title(doc):
    doc.add_page_break()
    scan_title = doc.add_paragraph().add_run('Функциональные показатели со сканера')
    scan_title.font.name = 'Times new roman'
    scan_title.font.size = Pt(14)
    scan_title.font.bold = True
    scan_title.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


# Полная "Сборка" 1 этажа для отчета
def add_floor_in_report(_path, doc, floor, floor_pictures_list, picture_number_in_file):
    add_floor_title_in_file(floor, doc)
    is_scanner_title_was_add = 0
    for picture in floor_pictures_list:
        script_number = get_script_number_of_picture(picture)
        if script_number in SCANNER_SCRIPTS_NUMBERS and is_scanner_title_was_add == 0:
            add_scanner_title(doc)
            is_scanner_title_was_add = 1
        add_picture_in_file(_path, doc, picture)
        add_picture_caption(picture, doc, picture_number_in_file)
        picture_number_in_file += 1
    doc.add_page_break()
    return picture_number_in_file
